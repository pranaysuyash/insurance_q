#!/usr/bin/env python3
"""Evaluate deterministic document capabilities from a versioned manifest.

The evaluator is deliberately local and source-safe. It measures native PDF
text/layout/table/embedded-image observations today and reports the remaining
OCR/VLM/formula/managed-provider gates as pending. It does not send documents
to a provider or persist source text unless ``--include-text`` is explicitly
passed for a synthetic fixture.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import sys
import time
from pathlib import Path
from typing import Any

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

import fitz  # noqa: E402  # repo-root bootstrap must precede local imports
from PIL import Image  # noqa: E402  # repo-root bootstrap must precede local imports
from docx import Document  # noqa: E402  # native DOCX benchmark fixture

from src.models.document_intelligence import build_document_cir  # noqa: E402
from src.ocr.native_pdf import extract_native_pdf_nodes  # noqa: E402
from src.ocr.native_office import extract_native_docx_document  # noqa: E402
from src.ocr.native_formats import (  # noqa: E402
    extract_native_eml_document,
    extract_native_html_document,
)
from src.utils.native_runtime import configure_native_library_paths  # noqa: E402


DEFAULT_MANIFEST = Path("docs/eval/document_intelligence/capability_manifest_v1.json")


def _synthetic_mixed_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page(width=320, height=240)
    page.insert_text((30, 30), "Health policy schedule")
    for x in (30, 150, 290):
        page.draw_line((x, 70), (x, 150))
    for y in (70, 110, 150):
        page.draw_line((30, y), (290, y))
    page.insert_text((40, 95), "Benefit")
    page.insert_text((160, 95), "Limit")
    page.insert_text((40, 135), "Health")
    page.insert_text((160, 135), "500000")
    image = Image.new("RGB", (20, 20), "#2f80ed")
    image_buffer = io.BytesIO()
    image.save(image_buffer, format="PNG")
    page.insert_image(fitz.Rect(30, 170, 90, 230), stream=image_buffer.getvalue())
    result = document.tobytes()
    document.close()
    return result


def _synthetic_scan() -> bytes:
    image = Image.new("RGB", (1000, 300), "white")
    from PIL import ImageDraw
    ImageDraw.Draw(image).text(
        (40, 80),
        "Health Policy 500000 POLICY-TEST-001",
        fill="black",
    )
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _synthetic_mixed_scan_pdf() -> bytes:
    scan_source = fitz.open()
    scan_source_page = scan_source.new_page(width=1000, height=300)
    scan_source_page.insert_text(
        (40, 120),
        "Health Policy 500000 POLICY-TEST-001",
        fontsize=44,
    )
    scan_image = scan_source_page.get_pixmap(dpi=72).tobytes("png")
    scan_source.close()

    document = fitz.open()
    native_page = document.new_page(width=1000, height=300)
    native_page.insert_text((40, 60), "Native policy schedule NATIVE-001")
    scan_page = document.new_page(width=1000, height=300)
    scan_page.insert_image(
        fitz.Rect(0, 0, 1000, 300),
        stream=scan_image,
    )
    result = document.tobytes()
    document.close()
    return result


def _synthetic_form_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page(width=320, height=240)
    page.insert_text((30, 30), "Policy number")
    widget = fitz.Widget()
    widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
    widget.field_name = "policy_number"
    widget.field_value = "POL-123"
    widget.rect = fitz.Rect(30, 40, 180, 60)
    page.add_widget(widget)
    result = document.tobytes()
    document.close()
    return result


def _synthetic_docx() -> bytes:
    document = Document()
    document.add_heading("Health Policy", level=1)
    document.add_paragraph("Coverage is subject to the policy schedule.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Benefit"
    table.cell(0, 1).text = "Limit"
    table.cell(1, 0).text = "Hospitalization"
    table.cell(1, 1).text = "500000"
    image = Image.new("RGB", (16, 16), "#2f80ed")
    image_buffer = io.BytesIO()
    image.save(image_buffer, format="PNG")
    document.add_picture(io.BytesIO(image_buffer.getvalue()))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _synthetic_html() -> bytes:
    return b"""
    <html><body><h1>Health Policy</h1>
    <p>Coverage is subject to the schedule.</p>
    <table><tr><th>Benefit</th><th>Limit</th></tr>
    <tr><td>Hospitalization</td><td>500000</td></tr></table>
    <img src="https://example.test/chart.png" alt="policy chart">
    </body></html>
    """


def _synthetic_eml() -> bytes:
    return (
        b"From: support@example.test\nTo: user@example.test\n"
        b"Subject: Policy schedule\nMIME-Version: 1.0\n"
        b"Content-Type: multipart/mixed; boundary=part\n\n"
        b"--part\nContent-Type: text/plain; charset=utf-8\n\n"
        b"Please review policy POL-123.\n"
        b"--part\nContent-Type: image/png\n"
        b"Content-Disposition: attachment; filename=chart.png\n"
        b"Content-Transfer-Encoding: base64\n\n"
        b"iVBORw0KGgo=\n--part--\n"
    )


def _normalize(text: str) -> str:
    return "".join(text.casefold().split())


def _evaluate_case(
    case: dict[str, Any],
    repo_root: Path,
    include_text: bool,
    ocr_profile: str | None = None,
    ocr_pipeline: Any = None,
) -> dict[str, Any]:
    started = time.monotonic()
    if case["kind"] == "fixture":
        fixture_path = repo_root / case["fixture"]
        if not fixture_path.is_file():
            raise FileNotFoundError(fixture_path)
        pdf_bytes = fixture_path.read_bytes()
        fixture_name = str(case["fixture"])
    elif case["kind"] == "generated_mixed_pdf":
        pdf_bytes = _synthetic_mixed_pdf()
        fixture_name = "generated:synthetic_mixed_native"
    elif case["kind"] == "generated_scan":
        if ocr_profile != case.get("profile"):
            return {
                "id": case["id"],
                "fixture": "generated:synthetic_scan",
                "status": "not_run",
                "required_profile": case.get("profile"),
                "reason": "run with --ocr-profile doctr to execute the local OCR gate",
            }
        pdf_bytes = _synthetic_scan()
        fixture_name = "generated:synthetic_scan"
        if ocr_pipeline is None:
            from src.ocr.pipeline import OCRPipeline
            ocr_pipeline = OCRPipeline()
        import asyncio
        ocr_text = asyncio.run(
            ocr_pipeline._get_ocr_text_for_image(pdf_bytes, page_num=1)
        )
        page_texts = {1: ocr_text}
        cir = build_document_cir(
            filename=fixture_name,
            source_artifact_sha256=hashlib.sha256(pdf_bytes).hexdigest(),
            file_type="png",
            page_texts=page_texts,
            page_images={1: pdf_bytes},
            parser_profile="local_doctr_ocr",
        )
        full_text = ocr_text
        normalized_text = _normalize(full_text)
        token_matches = {
            token: _normalize(token) in normalized_text
            for token in case.get("expected_tokens", [])
        }
        required = set(case.get("required_capabilities", []))
        observed = {capability.value for capability in cir.capabilities}
        result = {
            "id": case["id"],
            "fixture": fixture_name,
            "fixture_sha256": cir.source_artifact_sha256,
            "status": "passed" if required <= observed and all(token_matches.values()) else "failed",
            "elapsed_seconds": round(time.monotonic() - started, 4),
            "page_count": 1,
            "observed_capabilities": sorted(observed),
            "required_capabilities": sorted(required),
            "missing_capabilities": sorted(required - observed),
            "expected_token_matches": token_matches,
            "node_counts": {
                node_type: sum(node.node_type == node_type for node in cir.nodes)
                for node_type in sorted({node.node_type for node in cir.nodes})
            },
        }
        if include_text:
            result["source_text"] = full_text
        return result
    elif case["kind"] == "generated_mixed_scan_pdf":
        if ocr_profile != case.get("profile"):
            return {
                "id": case["id"],
                "fixture": "generated:synthetic_mixed_scan_pdf",
                "status": "not_run",
                "required_profile": case.get("profile"),
                "reason": "run with --ocr-profile doctr to execute the local mixed-page OCR gate",
            }
        pdf_bytes = _synthetic_mixed_scan_pdf()
        fixture_name = "generated:synthetic_mixed_scan_pdf"
        if ocr_pipeline is None:
            from src.ocr.pipeline import OCRPipeline
            ocr_pipeline = OCRPipeline()
        document = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            page_texts = {}
            page_images = {}
            for page_number in range(1, document.page_count + 1):
                page = document.load_page(page_number - 1)
                text = page.get_text().strip()
                if text:
                    page_texts[page_number] = text
                page_images[page_number] = page.get_pixmap(dpi=150).tobytes("png")
        finally:
            document.close()
        native_pages = set(page_texts)
        ocr_pages = set()
        unresolved_pages = []
        import asyncio
        for page_number in sorted(set(page_images) - native_pages):
            ocr_text = asyncio.run(
                ocr_pipeline._get_ocr_text_for_image(
                    page_images[page_number], page_num=page_number
                )
            ).strip()
            if ocr_text:
                page_texts[page_number] = ocr_text
                ocr_pages.add(page_number)
            else:
                unresolved_pages.append(page_number)
        nodes = extract_native_pdf_nodes(pdf_bytes)
        cir = build_document_cir(
            filename=fixture_name,
            source_artifact_sha256=hashlib.sha256(pdf_bytes).hexdigest(),
            file_type="pdf",
            page_texts=page_texts,
            page_images=page_images,
            parser_profile="mixed_native_doctr_ocr",
            observed_nodes=nodes,
            native_text_pages=native_pages,
            ocr_pages=ocr_pages,
        )
        full_text = "\n".join(page_texts[page] for page in sorted(page_texts))
        normalized_text = _normalize(full_text)
        token_matches = {
            token: _normalize(token) in normalized_text
            for token in case.get("expected_tokens", [])
        }
        required = set(case.get("required_capabilities", []))
        observed = {capability.value for capability in cir.capabilities}
        result = {
            "id": case["id"],
            "fixture": fixture_name,
            "fixture_sha256": cir.source_artifact_sha256,
            "status": "passed"
            if not unresolved_pages
            and required <= observed
            and all(token_matches.values())
            else "failed",
            "elapsed_seconds": round(time.monotonic() - started, 4),
            "page_count": 2,
            "observed_capabilities": sorted(observed),
            "required_capabilities": sorted(required),
            "missing_capabilities": sorted(required - observed),
            "expected_token_matches": token_matches,
            "native_text_pages": sorted(native_pages),
            "ocr_pages": sorted(ocr_pages),
            "unresolved_pages": unresolved_pages,
            "node_counts": {
                node_type: sum(node.node_type == node_type for node in cir.nodes)
                for node_type in sorted({node.node_type for node in cir.nodes})
            },
        }
        if include_text:
            result["source_text"] = full_text
        return result
    elif case["kind"] == "generated_docx_native":
        docx_bytes = _synthetic_docx()
        fixture_name = "generated:synthetic_docx_native"
        native_docx = extract_native_docx_document(docx_bytes)
        cir = build_document_cir(
            filename=fixture_name,
            source_artifact_sha256=native_docx["source_artifact_sha256"],
            file_type="docx",
            page_texts=native_docx["page_texts"],
            parser_profile=native_docx["parser_profile"],
            observed_nodes=native_docx["nodes"],
        )
        full_text = native_docx["full_text"]
        normalized_text = _normalize(full_text)
        token_matches = {
            token: _normalize(token) in normalized_text
            for token in case.get("expected_tokens", [])
        }
        required = set(case.get("required_capabilities", []))
        observed = {capability.value for capability in cir.capabilities}
        node_types = {node.node_type for node in cir.nodes}
        required_node_types = set(case.get("required_node_types", []))
        result = {
            "id": case["id"],
            "fixture": fixture_name,
            "fixture_sha256": cir.source_artifact_sha256,
            "status": "passed"
            if required <= observed
            and required_node_types <= node_types
            and all(token_matches.values())
            else "failed",
            "elapsed_seconds": round(time.monotonic() - started, 4),
            "page_count": 1,
            "observed_capabilities": sorted(observed),
            "required_capabilities": sorted(required),
            "missing_capabilities": sorted(required - observed),
            "observed_node_types": sorted(node_types),
            "missing_node_types": sorted(required_node_types - node_types),
            "expected_token_matches": token_matches,
            "node_counts": {
                node_type: sum(node.node_type == node_type for node in cir.nodes)
                for node_type in sorted(node_types)
            },
        }
        if include_text:
            result["source_text"] = full_text
        return result
    elif case["kind"] in {"generated_html_native", "generated_eml_native"}:
        if case["kind"] == "generated_html_native":
            source_bytes = _synthetic_html()
            native_document = extract_native_html_document(source_bytes)
            fixture_name = "generated:synthetic_html_native"
        else:
            source_bytes = _synthetic_eml()
            native_document = extract_native_eml_document(source_bytes)
            fixture_name = "generated:synthetic_eml_native"
        cir = build_document_cir(
            filename=fixture_name,
            source_artifact_sha256=native_document["source_artifact_sha256"],
            file_type="html" if case["kind"] == "generated_html_native" else "eml",
            page_texts=native_document["page_texts"],
            parser_profile=native_document["parser_profile"],
            observed_nodes=native_document["nodes"],
        )
        full_text = native_document["full_text"]
        normalized_text = _normalize(full_text)
        token_matches = {
            token: _normalize(token) in normalized_text
            for token in case.get("expected_tokens", [])
        }
        required = set(case.get("required_capabilities", []))
        observed = {capability.value for capability in cir.capabilities}
        node_types = {node.node_type for node in cir.nodes}
        required_node_types = set(case.get("required_node_types", []))
        result = {
            "id": case["id"],
            "fixture": fixture_name,
            "fixture_sha256": cir.source_artifact_sha256,
            "status": "passed"
            if required <= observed
            and required_node_types <= node_types
            and all(token_matches.values())
            else "failed",
            "elapsed_seconds": round(time.monotonic() - started, 4),
            "page_count": 1,
            "observed_capabilities": sorted(observed),
            "required_capabilities": sorted(required),
            "missing_capabilities": sorted(required - observed),
            "observed_node_types": sorted(node_types),
            "missing_node_types": sorted(required_node_types - node_types),
            "expected_token_matches": token_matches,
            "node_counts": {
                node_type: sum(node.node_type == node_type for node in cir.nodes)
                for node_type in sorted(node_types)
            },
        }
        if include_text:
            result["source_text"] = full_text
        return result
    elif case["kind"] == "generated_form_pdf":
        pdf_bytes = _synthetic_form_pdf()
        fixture_name = "generated:synthetic_form_pdf"
    else:
        raise ValueError(f"Unsupported benchmark case kind: {case['kind']}")

    document = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        page_texts = {
            page_number: document.load_page(page_number - 1).get_text()
            for page_number in range(1, document.page_count + 1)
        }
    finally:
        document.close()

    nodes = extract_native_pdf_nodes(pdf_bytes)
    cir = build_document_cir(
        filename=fixture_name,
        source_artifact_sha256=hashlib.sha256(pdf_bytes).hexdigest(),
        file_type="pdf",
        page_texts=page_texts,
        parser_profile="pymupdf_native",
        observed_nodes=nodes,
    )
    full_text = "\n".join(page_texts.values())
    normalized_text = _normalize(full_text)
    token_matches = {
        token: _normalize(token) in normalized_text
        for token in case.get("expected_tokens", [])
    }
    observed = {capability.value for capability in cir.capabilities}
    required = set(case.get("required_capabilities", []))
    result: dict[str, Any] = {
        "id": case["id"],
        "fixture": fixture_name,
        "fixture_sha256": cir.source_artifact_sha256,
        "status": "passed" if required <= observed and all(token_matches.values()) else "failed",
        "elapsed_seconds": round(time.monotonic() - started, 4),
        "page_count": len(page_texts),
        "observed_capabilities": sorted(observed),
        "required_capabilities": sorted(required),
        "missing_capabilities": sorted(required - observed),
        "expected_token_matches": token_matches,
        "node_counts": {
            node_type: sum(node.node_type == node_type for node in cir.nodes)
            for node_type in sorted({node.node_type for node in cir.nodes})
        },
    }
    if include_text:
        result["source_text"] = full_text
    return result


def evaluate_manifest(
    manifest_path: Path,
    *,
    repo_root: Path,
    include_text: bool = False,
    ocr_profile: str | None = None,
) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    ocr_pipeline = None
    if ocr_profile == "doctr":
        # Keep the standalone benchmark on the same native-library contract as
        # the API entrypoint.  On macOS, doctr imports WeasyPrint/Pango/GLib
        # transitively; the packages can be installed in the project venv
        # while the dynamic loader still needs the Homebrew prefixes.
        configure_native_library_paths()
        from src.ocr.pipeline import OCRPipeline
        ocr_pipeline = OCRPipeline()
    case_results = [
        _evaluate_case(case, repo_root, include_text, ocr_profile, ocr_pipeline)
        for case in manifest.get("cases", [])
    ]
    executed = [result for result in case_results if result["status"] != "not_run"]
    return {
        "manifest_id": manifest["manifest_id"],
        "schema_version": manifest["schema_version"],
        "manifest_sha256": hashlib.sha256(manifest_path.read_bytes()).hexdigest(),
        "case_results": case_results,
        "capability_gates": manifest.get("capability_gates", []),
        "all_executed_cases_passed": all(result["status"] == "passed" for result in executed),
        "not_run_cases": sum(result["status"] == "not_run" for result in case_results),
        "ocr_profile": ocr_profile,
        "source_text_included": include_text,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--include-text", action="store_true")
    parser.add_argument("--ocr-profile", choices=["doctr"], help="Run an optional local OCR gate")
    parser.add_argument("--strict", action="store_true", help="Exit 2 if an executed case fails")
    args = parser.parse_args()
    repo_root = REPO_ROOT
    report = evaluate_manifest(
        args.manifest.resolve(),
        repo_root=repo_root,
        include_text=args.include_text,
        ocr_profile=args.ocr_profile,
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, indent=2))
    return 2 if args.strict and (
        not report["all_executed_cases_passed"] or report["not_run_cases"]
    ) else 0


if __name__ == "__main__":
    raise SystemExit(main())
